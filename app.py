import io
import os
import re
import tempfile
from datetime import datetime
import pandas as pd
import streamlit as st
import requests
from docx import Document
from docx.shared import Inches, Pt, Cm

st.set_page_config(
    page_title="Gerador de Fichas", page_icon="📄", layout="centered"
)

st.title("📄 Gerador de Fichas (Busca Direta)")

st.markdown("""
### 📢 Como usar as Fotos pelo Google Drive:
1. Deixe todas as fotos direto na pasta do **Google Drive** (sem subpastas).
2. Garanta que o nome de cada arquivo seja o nome exato do servo (ex: `Nome do Servo.jpg` ou `Nome do Servo.png`).
3. Compartilhe a pasta como **"Qualquer pessoa com o link"** (Leitor).
""")

logo_file = st.file_uploader(
    "1. Logo do Acampamento (opcional)", type=["png", "jpg", "jpeg"]
)

csv_file = st.file_uploader("2. Arquivo CSV das Inscrições", type=["csv"])

pasta_drive_url = st.text_input(
    "3. Cole aqui o Link da Pasta do Google Drive contendo as fotos"
)

titulo_acampamento = st.text_input(
    "Título do Acampamento", value="Ficha de Inscrição do Servo - Resumida"
)


def calcular_idade(data_nascimento):
    try:
        nascimento = pd.to_datetime(
            data_nascimento, dayfirst=True, errors="coerce"
        )
        if pd.isna(nascimento):
            return ""
        hoje = datetime.today()
        idade = hoje.year - nascimento.year
        if (hoje.month, hoje.day) < (nascimento.month, nascimento.day):
            idade -= 1
        return idade
    except:
        return ""


def extrair_id_pasta(url):
    match = re.search(r"folders/([a-zA-Z0-9-_]+)", url)
    if match:
        return match.group(1)
    return None


# Função que tenta buscar e indexar varrendo a estrutura pública de visualização do Drive
def obter_arquivos_da_pasta_drive(folder_id):
    dict_fotos = {}
    try:
        url_alt = f"https://drive.google.com/embeddedfolderview?id={folder_id}"
        res_alt = requests.get(url_alt, timeout=15)
        # Regex aprimorado para capturar nomes e IDs direto do HTML estruturado do Drive
        matches_alt = re.findall(r'id="entry-([a-zA-Z0-9-_]+)".*?class="cl-name"[^>]*>([^<]+)', res_alt.text, re.DOTALL)
        for file_id, name in matches_alt:
            nome_sem_extensao = os.path.splitext(name.strip())[0].strip().lower()
            dict_fotos[nome_sem_extensao] = file_id
    except:
        pass
    return dict_fotos


if "word_file" not in st.session_state:
    st.session_state.word_file = None

if csv_file:
    try:
        try:
            df = pd.read_csv(csv_file, sep=";", encoding="utf-8-sig")
        except:
            csv_file.seek(0)
            df = pd.read_csv(csv_file, sep=";", encoding="latin1")

        df = df.dropna(how='all')
        st.success(f"{len(df)} inscrições válidas carregadas.")

        dicionario_fotos_drive = {}
        id_pasta = extrair_id_pasta(pasta_drive_url) if pasta_drive_url else None

        if id_pasta:
            with st.spinner("Conectando à pasta do Google Drive..."):
                dicionario_fotos_drive = obter_arquivos_da_pasta_drive(id_pasta)
            
            if dicionario_fotos_drive:
                st.info(f"Conectado! {len(dicionario_fotos_drive)} fotos pré-mapeadas.")
            else:
                st.warning("⚠️ O Google limitou a listagem automática prévia. O sistema fará a busca individual direta por nome de cada servo ao clicar no botão.")

        if st.button("Gerar Fichas"):
            fotos_com_erro = []
            
            try:
                with st.spinner("Processando fichas... Buscando e baixando imagens."):
                    doc = Document()
                    
                    sections = doc.sections
                    for section in sections:
                        section.top_margin = Cm(1.5)
                        section.bottom_margin = Cm(1.5)
                        section.left_margin = Cm(1.5)
                        section.right_margin = Cm(1.5)

                    logo_temp = None
                    if logo_file:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                            tmp.write(logo_file.getvalue())
                            logo_temp = tmp.name

                    col_nome = "Nome"
                    col_cidade = "Cidade"
                    col_telefone = "Celular"
                    col_nascimento = "Data de Nascimento"
                    col_camiseta = "Qual o tamanho da sua camiseta?"
                    col_ministerios = (
                        "A coordenação do acampamento é a responsável por montar as equipes de trabalho. "
                        "Mas, gostaríamos de receber a sua opinião. Assinale até 3 ministérios que você gostaria de servir. "
                    )
                    col_serviu = "Já serviu em acampamentos? Se sim, quais ministérios?"
                    col_pastoral = "Participa de alguma Pastoral ou Movimento? Se sim, qual:"
                    col_sacramentos = "Quais Sacramentos possui (batismo, eucaristia, crisma, matrimônio e ordem)?"

                    for i, row in df.iterrows():
                        if pd.isna(row.get(col_nome)) or str(row.get(col_nome)).strip() == "":
                            continue

                        nome = str(row.get(col_nome, "")).strip()

                        if i > 0:
                            doc.add_page_break()

                        if logo_temp:
                            p_logo = doc.add_paragraph()
                            p_logo.alignment = 1
                            run_logo = p_logo.add_run()
                            run_logo.add_picture(logo_temp, width=Inches(1.8))

                        titulo = doc.add_paragraph()
                        titulo.alignment = 1
                        run_titulo = titulo.add_run(titulo_acampamento)
                        run_titulo.bold = True
                        run_titulo.font.size = Pt(13)

                        cidade = row.get(col_cidade, "") if pd.notna(row.get(col_cidade)) else ""
                        telefone = row.get(col_telefone, "") if pd.notna(row.get(col_telefone)) else ""
                        nascimento = row.get(col_nascimento, "") if pd.notna(row.get(col_nascimento)) else ""
                        camiseta = row.get(col_camiseta, "") if pd.notna(row.get(col_camiseta)) else ""
                        ministerios = row.get(col_ministerios, "") if pd.notna(row.get(col_ministerios)) else ""
                        serviu = row.get(col_serviu, "") if pd.notna(row.get(col_serviu)) else ""
                        pastoral = row.get(col_pastoral, "") if pd.notna(row.get(col_pastoral)) else ""
                        sacramentos = row.get(col_sacramentos, "") if pd.notna(row.get(col_sacramentos)) else ""

                        idade = calcular_idade(nascimento)

                        nome_p = doc.add_paragraph()
                        nome_p.alignment = 1
                        run_nome = nome_p.add_run(nome.upper())
                        run_nome.bold = True
                        run_nome.font.size = Pt(24)

                        doc.add_paragraph("")

                        tabela = doc.add_table(rows=0, cols=2)
                        tabela.style = "Table Grid"

                        if idade != "":
                            nascimento_texto = f"{nascimento} ({idade} anos)"
                        else:
                            nascimento_texto = str(nascimento)

                        campos = [
                            ("Cidade", cidade),
                            ("Telefone", telefone),
                            ("Nascimento / Idade", nascimento_texto),
                            ("Tamanho Camiseta", camiseta),
                            ("Ministérios Desejados", ministerios),
                            ("Já Serviu", serviu),
                            ("Pastoral / Movimento", pastoral),
                            ("Sacramentos", sacramentos),
                        ]

                        for campo, valor in campos:
                            linha = tabela.add_row().cells
                            p_campo = linha[0].paragraphs[0]
                            run_campo = p_campo.add_run(str(campo))
                            run_campo.bold = True
                            run_campo.font.size = Pt(11)
                            
                            p_valor = linha[1].paragraphs[0]
                            run_valor = p_valor.add_run(str(valor))
                            run_valor.font.size = Pt(11)

                        p_espaco = doc.add_paragraph()
                        p_espaco.paragraph_format.space_before = Pt(8)

                        foto = doc.add_table(rows=1, cols=1)
                        foto.style = "Table Grid"
                        celula = foto.cell(0, 0)
                        p_foto = celula.paragraphs[0]
                        p_foto.alignment = 1

                        nome_chave = nome.lower()
                        id_foto_drive = dicionario_fotos_drive.get(nome_chave)
                        
                        # ESTRATÉGIA DE BUSCA DINÂMICA CASO A LISTAGEM DO GOOGLE TENHA FALHADO
                        id_encontrado = None
                        if id_foto_drive:
                            id_encontrado = id_foto_drive
                        elif id_pasta:
                            # Se a listagem falhou, tentamos forçar a varredura da foto individual pelo nome exato do servo via Web View
                            try:
                                url_busca = f"https://drive.google.com/embeddedfolderview?id={id_pasta}&q={nome_chave}"
                                res_busca = requests.get(url_busca, timeout=5)
                                match_id = re.search(r'id="entry-([a-zA-Z0-9-_]+)"', res_busca.text)
                                if match_id:
                                    id_encontrado = match_id.group(1)
                            except:
                                pass

                        if id_encontrado:
                            try:
                                url_download = f"https://docs.google.com/uc?export=download&id={id_encontrado}"
                                session = requests.Session()
                                response_foto = session.get(url_download, stream=True, timeout=15)
                                
                                token = None
                                for key, val in response_foto.cookies.items():
                                    if key.startswith('download_warning'):
                                        token = val
                                if token:
                                    url_download += f"&confirm={token}"
                                    response_foto = session.get(url_download, stream=True, timeout=15)

                                if response_foto.status_code == 200:
                                    # Usa o formato correto de imagem com base no content-type retornado
                                    content_type = response_foto.headers.get('content-type', '')
                                    suffix = ".png" if "png" in content_type else ".jpg"
                                    
                                    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as f_tmp:
                                        f_tmp.write(response_foto.content)
                                        foto_temp_path = f_tmp.name
                                    
                                    run_foto = p_foto.add_run()
                                    run_foto.add_picture(foto_temp_path, width=Inches(1.6))
                                    
                                    if os.path.exists(foto_temp_path):
                                        os.remove(foto_temp_path)
                                else:
                                    raise Exception()
                            except Exception:
                                fotos_com_erro.append(nome.upper())
                                run_foto = p_foto.add_run("\n\n\n\nERRO AO BAIXAR IMAGEM\n\n\n\n")
                                run_foto.font.size = Pt(11)
                        else:
                            run_foto = p_foto.add_run("\n\n\n\nFOTO NÃO ENCONTRADA\n\n\n\n")
                            run_foto.font.size = Pt(11)

                        p_linha = doc.add_paragraph()
                        p_linha.alignment = 1
                        p_linha.paragraph_format.space_before = Pt(8)
                        p_linha.add_run("________________________________________")

                    bio = io.BytesIO()
                    doc.save(bio)
                    bio.seek(0)

                    st.session_state.word_file = bio.read()

                    if logo_temp and os.path.exists(logo_temp):
                        os.remove(logo_temp)
                        
                st.success("✨ Processamento concluído com sucesso!")
                
                if fotos_com_erro:
                    st.warning(f"⚠️ As fichas foram geradas, mas {len(fotos_com_erro)} fotos falharam no download:")
                    for servo_errado in fotos_com_erro:
                        st.write(f"- {servo_errado}")

            except Exception as e:
                st.error(f"Erro crítico no loop de geração: {e}")

        if st.session_state.word_file is not None:
            st.download_button(
                label="📥 Baixar Arquivo Word",
                data=st.session_state.word_file,
                file_name="fichas_acampamento.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    except Exception as e:
        st.error(f"Erro geral no processamento: {e}")
